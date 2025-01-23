# --- modules/file_handler.py ---
from pathlib import Path
import docx

def load_word_documents(directory: Path) -> dict[str, list[str]]:
    """Charge tous les documents Word d'un répertoire et extrait leur texte."""
    documents = {}
    for file in directory.glob("*.docx"):
        doc = docx.Document(file)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        documents[file.name] = text
    return documents

def save_anonymized_documents(directory: Path, documents: dict[str, list[str]]) -> None:
    """Sauvegarde les documents Word anonymisés."""
    anonymized_dir = directory / "anonymized"
    anonymized_dir.mkdir(exist_ok=True)
    
    for filename, content in documents.items():
        doc = docx.Document()
        for line in content:
            doc.add_paragraph(line)
        doc.save(anonymized_dir / filename)
        
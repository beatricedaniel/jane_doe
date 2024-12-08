import os
import re
import csv 
import json
from utils.load_docx_files import load_docx_files
from utils.get_filename import get_filename
from docx import Document

class Anonymizer():

    def get_patterns(settings: dict) -> list:
        """Get list of patterns from the config dictionary.

        Args:
            settings: dictionary.

        Returns:
            list of patterns.
        """
        # Retrieve the list of patterns from the settings
        patterns = settings.get("anonymization_patterns", [])
        try:
            return patterns
        except ValueError:
            raise ValueError("Patterns must be a list")
        
    def get_sensitive_infos(patterns: list, input_dir_path: str, output_csv_path: str) -> list:
        """Get list of words corresponding to the patterns.

        Args:
            patterns: list of regex patterns.
            input_dir_path: directory where are the input documents.
            output_csv_path: path of the csv file where are stored results.

        Returns:
            csv file with 
                a acolumn for the sensitive infos extracted
                a column for the location of each word.
        """
        # Compile the regex patterns
        compiled_patterns = [(one_pattern, re.compile(one_pattern)) for one_pattern in patterns]

        # Retrieve list of docx
        docx_directory = load_docx_files(input_dir_path)

        # Loop to extract and store sensitive informations from each docx
        for docx_path in docx_directory:
            # Convert file into format exploitable with docx library
            document = Document(docx_path)
            results = []

            # Extract from paragraphs
            for i, paragraph in enumerate(document.paragraphs):
                for pattern_text, pattern in compiled_patterns:
                    matches = pattern.findall(paragraph.text)
                    for match in matches:
                        results.append({
                            "word": match,
                            "regex_pattern": pattern_text
                        })

            # Extract from tables
            for table_idx, table in enumerate(document.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        for pattern_text, pattern in compiled_patterns:
                            matches = pattern.findall(cell.text)
                            for match in matches:
                                results.append({
                                    "word": match,
                                    "regex_pattern": pattern_text
                                })

            # Retrieve name of each input file
            input_file_name = get_filename(docx_path)

            # Replace the '*' in the path with the file name
            output_file_path = output_csv_path.replace("*", input_file_name)
            
            # Store results info a .csv file
            with open(output_file_path, mode='w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['word', 'regex_pattern']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                # Write header
                writer.writeheader()
                
                # Write rows
                for result in results:
                    writer.writerow(result)

            print(f"Results saved to {output_file_path}")
            
        return
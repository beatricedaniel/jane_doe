import json
import csv
import re
from pathlib import Path
from presidio_analyzer import AnalyzerEngine
from docx import Document

class WordDocumentAnonymizer:
    def __init__(self, config_file: str):
        """
        Initialize the anonymizer with configuration data from a JSON file.
        
        :param config_file: Path to the JSON file containing regex patterns and variables.
        """
        self.config = self.load_config(config_file)
        self.regex_patterns = self.config.get("regex_patterns", [])
        self.output_csv_dir = self.config.get("output_csv_dir", "matches.csv")
        self.input_dir = self.config.get("input_dir", "input_documents")
        self.output_dir = self.config.get("output_dir", "anonymized_documents")
        self.analyzer = AnalyzerEngine()

    @staticmethod
    def load_config(config_file: str) -> dict:
        """
        Load configuration data from a JSON file.
        
        :param config_file: Path to the JSON file containing configuration data.
        :return: Dictionary containing configuration data.
        """
        with open(config_file, 'r') as file:
            return json.load(file)

    def find_matches(self, text: str) -> list:
        """
        Find matches in the text based on the loaded regex patterns.
        
        :param text: Text to search for matches.
        :return: List of tuples with the match, regex pattern, and NER analysis.
        """
        matches = []
        for pattern in self.regex_patterns:
            for match in re.finditer(pattern, text):
                match_text = match.group()
                # Analyze with Presidio
                presidio_results = self.analyzer.analyze(text=match_text, language="en")
                is_ner = "Yes" if presidio_results else "No"
                entity_type = presidio_results[0].entity_type if presidio_results else "N/A"
                matches.append((match_text, pattern, is_ner, entity_type))
        return matches

    def anonymize_text(self, text: str, matches: list) -> str:
        """
        Anonymize text by replacing NER matches with placeholders.
        
        :param text: Original text.
        :param matches: List of matches.
        :return: Anonymized text.
        """
        for match_text, _, is_ner, entity_type in matches:
            if is_ner == "Yes":
                placeholder = f"<{entity_type}>"
                text = text.replace(match_text, placeholder)
        return text

    def process_document(self, input_docx: str, output_docx: str, output_csv: str):
        """
        Process a Word document for anonymization.
        
        :param input_docx: Path to the input Word document.
        :param output_docx: Path to the output anonymized Word document.
        """
        doc = Document(input_docx)
        matches = []

        # Process paragraphs
        for paragraph in doc.paragraphs:
            paragraph_matches = self.find_matches(paragraph.text)
            matches.extend(paragraph_matches)
            paragraph.text = self.anonymize_text(paragraph.text, paragraph_matches)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_matches = self.find_matches(cell.text)
                    matches.extend(cell_matches)
                    cell.text = self.anonymize_text(cell.text, cell_matches)

        # Save the anonymized document
        doc.save(output_docx)

        # Write matches to CSV
        self.write_matches_to_csv(output_csv, matches)

    def write_matches_to_csv(self, output_csv: str, matches: list):
        """
        Write matches to a CSV file.
        
        :param output_csv: Path to the output CSV file.
        :param matches: List of matches.
        """
        with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(["Matched Text", "Regex Pattern", "Is NER", "Entity Type"])
            writer.writerows(matches)
        print(f"CSV file written to: {output_csv}")

    def resolve_csv_path(self, input_file: str) -> str:
        """
        Generate the CSV path by replacing '*' in the configured output_csv_dir with the input file name (without extension).
        
        :param input_file: Path to the input file.
        :return: Resolved path to the output CSV file.
        """
        input_name = Path(input_file).stem  # Get file name without extension
        return self.output_csv_dir.replace("*", input_name)

import os
import spacy
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig, RecognizerResult
from presidio_analyzer.nlp_engine import SpacyNlpEngine
from presidio_analyzer import RecognizerResult, PatternRecognizer
from docx import Document
import csv

# Configuration : chemins des répertoires et fichier CSV
INPUT_DIR = "../../../data/input"
OUTPUT_DIR = "data_output"
CSV_PATH = "data_output/entities.csv"
PREFIX = "anonymized_"
REPLACEMENT_TEXT = "[ANONYMOUS]"

# Initialisation du moteur SpaCy en français
print("Chargement du modèle SpaCy...")
nlp = spacy.load("fr_core_news_md")
nlp_engine = SpacyNlpEngine()
nlp_engine.nlp = {"fr": nlp}

# Initialisation des moteurs Presidio
print("Initialisation de Presidio...")
# analyzer = AnalyzerEngine(nlp_engine=nlp_engine, default_language="fr")
analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
anonymizer = AnonymizerEngine()

# Configuration de reconnaissances personnalisées (ajoutez ici si besoin)
# custom_recognizer = PatternRecognizer(supported_entity="EMAIL_ADDRESS", patterns=[
#     {"name": "email_pattern", "pattern": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", "score": 0.8}
# ])
# analyzer.registry.add_recognizer(custom_recognizer)

# Création du fichier CSV pour enregistrer les entités détectées
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

with open(CSV_PATH, mode="w", newline="", encoding="utf-8") as csv_file:
    writer = csv.writer(csv_file)
    writer.writerow(["Document", "Entity"])

    # Traitement des documents Word
    for filename in os.listdir(INPUT_DIR):
        if filename.endswith(".docx"):
            input_path = os.path.join(INPUT_DIR, filename)
            output_path = os.path.join(OUTPUT_DIR, PREFIX + filename)

            print(f"Traitement de {filename}...")
            
            # Lecture du document Word
            doc = Document(input_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])

            # Analyse pour trouver les entités
            results = analyzer.analyze(
                text=full_text,
                language="fr",
                entities=[],  # Toutes les entités par défaut + personnalisées
                score_threshold=0.5
            )

            # Écriture des entités dans le CSV
            for result in results:
                # writer.writerow([filename, result.entity_type])
                writer.writerow([filename, full_text[result.start:result.end]])

            # Anonymisation du texte
            anonymized_results = anonymizer.anonymize(
                text=full_text,
                analyzer_results=results,
                # anonymizers_config={"DEFAULT": {"type": "replace", "new_value": REPLACEMENT_TEXT}}
                operators={"PERSON": OperatorConfig("replace", {"new_value": REPLACEMENT_TEXT})}
            )

            # Création du nouveau document Word anonymisé
            anonymized_doc = Document()
            for paragraph in anonymized_results.text.split("\n"):
                anonymized_doc.add_paragraph(paragraph)

            anonymized_doc.save(output_path)

print("Anonymisation terminée. Résultats enregistrés dans :")
print(f" - Dossier : {OUTPUT_DIR}")
print(f" - Fichier CSV : {CSV_PATH}")
